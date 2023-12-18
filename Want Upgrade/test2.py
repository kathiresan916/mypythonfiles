from docx import Document

# Create or open a Word document
doc = Document()

# Add content to the document
doc.add_paragraph("Modified")

# Save the document
doc.save("Test.docx")

# Read the content (just for demonstration purposes)
doc = Document("Test.docx")
content = [paragraph.text for paragraph in doc.paragraphs]
print("\n".join(content))
